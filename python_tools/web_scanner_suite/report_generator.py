from docx import Document
import json

report = Document()
report.add_heading("Informe Profesional de Seguridad Web", 0)
report.add_heading("Resumen Ejecutivo", level=1)
report.add_paragraph("Este informe automatizado describe hallazgos de seguridad web en entornos controlados.")

report.add_heading("Hallazgos", level=1)
for file, label in [("sqli_results.json", "SQL Injection"), ("xss_results.json", "Cross Site Scripting")]:
    try:
        data = json.load(open(file))
        report.add_heading(label, level=2)
        for finding in data:
            report.add_paragraph(f"- {finding['tipo']} en {finding['url']} usando {finding['payload']}")
    except:
        continue

report.add_heading("Recomendaciones", level=1)
report.add_paragraph("Validar y sanitizar todas las entradas. Usar ORM parametrizados. Aplicar CSP para prevenir XSS.")

report.add_heading("Reflexión Ética", level=1)
report.add_paragraph("Todas las pruebas fueron realizadas en entornos vulnerables destinados a prácticas éticas. Se respetaron principios de no daño, responsabilidad profesional y documentación clara de resultados.")

report.save("informe_final.docx")
print("[+] Informe generado como informe_final.docx")
